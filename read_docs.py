import docx

with open(r'C:\Users\Student\output_docs.txt', 'w', encoding='utf-8') as f:
    f.write("="*60 + "\n")
    f.write("PART1_SITE_CONTENT_PLAN\n")
    f.write("="*60 + "\n")
    doc = docx.Document(r'C:\Users\Student\Desktop\portfolio-website\public\Opnsense\Datei\Part1_Site_Content_Plan.docx')
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + "\n")

    f.write("\n" + "="*60 + "\n")
    f.write("PART2_ENGINEERING_DOCUMENTATION (Headings)\n")
    f.write("="*60 + "\n")
    doc2 = docx.Document(r'C:\Users\Student\Desktop\portfolio-website\public\Opnsense\Datei\Part2_Engineering_Documentation.docx')
    for para in doc2.paragraphs:
        if para.text.strip():
            style_name = para.style.name if hasattr(para.style, 'name') else 'unknown'
            if 'Heading' in style_name or 'heading' in style_name:
                f.write(f"[{style_name}] {para.text}\n")

print("Done")
